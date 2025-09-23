import re
import pymupdf
from docx import Document
import torch
from transformers import AutoTokenizer, AutoModelForTokenClassification
from transformers import pipeline


class IntelligentContractExtractor:
    def __init__(self):
        # 加载预训练模型（中文法律领域优化模型）
        # 模型会自动下载，首次运行需要联网
        self.model_name = "uie-base"  # 百度开源的通用信息抽取模型
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        self.model = AutoModelForTokenClassification.from_pretrained(self.model_name)

        # 定义需要提取的合同要素
        self.contract_elements = [
            "甲方", "乙方", "合同金额", "履行期限",
            "违约责任", "签订日期", "合同编号",
            "甲方签名", "乙方签名", "法定代表人"
        ]

        # 创建信息抽取管道
        self.extractor = pipeline(
            "token-classification",
            model=self.model,
            tokenizer=self.tokenizer,
            aggregation_strategy="average"
        )

    def read_word_file(self, file_path):
        """读取Word文件内容"""
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # 读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"读取Word文件出错: {str(e)}")
            return ""

    def read_pdf_file(self, file_path):
        """读取PDF文件内容"""
        try:
            doc = pymupdf.open(file_path)
            full_text = []
            for page in doc:
                full_text.append(page.get_text())
            return "\n".join(full_text)
        except Exception as e:
            print(f"读取PDF文件出错: {str(e)}")
            return ""

    def clean_text(self, text):
        """清理文本中的冗余信息"""
        # 去除多余空白和特殊字符
        text = re.sub(r'\s+', ' ', text)
        # 去除无意义的标点符号组合
        text = re.sub(r'[^\w\s，,。.；;：:（）()《》<>“”""\'\'\d]', '', text)
        return text.strip()

    def extract_elements(self, text):
        """使用NLP模型提取合同要素"""
        results = {}

        # 对每个要素进行定向抽取
        for element in self.contract_elements:
            # 构造抽取提示
            prompt = f"抽取文本中的{element}信息：{text[:500]}..."  # 限制输入长度
            try:
                # 使用模型抽取
                outputs = self.extractor(prompt)

                # 整理抽取结果
                if outputs:
                    # 取置信度最高的结果
                    best_result = max(outputs, key=lambda x: x['score'])
                    element_value = best_result['word']

                    # 过滤无效结果
                    if len(element_value) > 1 and best_result['score'] > 0.6:
                        results[element] = {
                            "value": element_value,
                            "confidence": round(best_result['score'], 2)
                        }
            except Exception as e:
                print(f"提取{element}时出错: {str(e)}")
                continue

        # 特殊处理：将履行期限拆分为开始和结束日期
        if "履行期限" in results:
            term = results["履行期限"]["value"]
            # 简单拆分（模型已能识别大部分格式）
            if "至" in term:
                start, end = term.split("至", 1)
                results["履行期限"]["value"] = {
                    "start_date": start.strip(),
                    "end_date": end.strip()
                }

        return results

    def process_file(self, file_path):
        """处理合同文件并返回提取结果"""
        # 根据文件类型读取内容
        if file_path.lower().endswith('.docx'):
            text = self.read_word_file(file_path)
        elif file_path.lower().endswith('.pdf'):
            text = self.read_pdf_file(file_path)
        else:
            print("不支持的文件格式，仅支持.docx和.pdf")
            return None

        if not text:
            return None

        # 清理文本并提取要素
        clean_text = self.clean_text(text)
        return self.extract_elements(clean_text)


if __name__ == "__main__":
    # 初始化提取器
    extractor = IntelligentContractExtractor()

    # 处理文件
    file_path = "111.docx"  # 替换为你的合同文件路径
    print(f"正在处理文件: {file_path}")
    result = extractor.process_file(file_path)

    # 显示结果
    if result:
        print("\n===== 合同关键要素提取结果 =====")
        for key, info in result.items():
            if key == "履行期限":
                print(f"{key}：自 {info['value']['start_date']} 至 {info['value']['end_date']}")
            else:
                print(f"{key}：{info['value']}")
            print(f"   置信度：{info['confidence']:.2f}\n")
    else:
        print("未能提取到有效信息")
