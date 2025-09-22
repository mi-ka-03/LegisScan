from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from paddleocr import PaddleOCR
import numpy as np
import cv2
from io import BytesIO, StringIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import os
import uuid

# ======================== 初始化 ========================
app = Flask(__name__)
CORS(app)

# 初始化PaddleOCR
# 建议使用PP-OCRv4版本，在精度和速度上都有提升
# PaddleOCR会自动下载模型到 ~/.paddleocr/ 目录
ocr = PaddleOCR(use_angle_cls=True, lang='ch')

# 用于存放临时生成的Word文件
UPLOAD_FOLDER = 'temp_docs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ======================== 核心功能函数 ========================
def create_word_document_from_ocr(ocr_result):
    """
    根据OCR结果创建Word文档。
    :param ocr_result: PaddleOCR返回的识别结果，格式为 [ [ [ [poly], [text, score] ], ... ] ]
    :return: 生成的Word文档的临时文件路径
    """
    doc = Document()

    # 添加标题
    doc.add_heading('AI识别仅供参考', 0)

    if not ocr_result or not ocr_result[0]:
        doc.add_paragraph("未识别到任何文字。")
        temp_path = os.path.join(UPLOAD_FOLDER, f"empty_{uuid.uuid4()}.docx")
        doc.save(temp_path)
        return temp_path

    # 提取文本块和其位置信息
    # 每个文本块是 (文本内容, 左上角x, 左上角y, 右下角x, 右下角y)
    text_blocks = []
    for line in ocr_result[0]:
        poly, (text, score) = line
        x1, y1 = poly[0]
        x2, y2 = poly[2]
        text_blocks.append((text, int(x1), int(y1), int(x2), int(y2)))

    # --- 布局还原核心逻辑 ---
    # 1. 根据y坐标对文本块进行分组，模拟行
    lines = []
    current_line = []
    # 假设行与行之间的垂直距离大于20像素则视为新行，可根据实际情况调整
    line_height_threshold = 20
    # 按y坐标排序
    text_blocks.sort(key=lambda b: b[2])

    for block in text_blocks:
        if not current_line:
            current_line.append(block)
        else:
            # 如果当前块的y坐标与行首块的y坐标差小于阈值，则认为在同一行
            if block[2] - current_line[0][2] < line_height_threshold:
                current_line.append(block)
            else:
                lines.append(current_line)
                current_line = [block]
    if current_line:
        lines.append(current_line)

    # 2. 逐行处理，写入Word文档
    for line in lines:
        # 同一行内，按x坐标从左到右排序
        line.sort(key=lambda b: b[1])

        # 将该行所有文本块合并成一个段落
        paragraph_text = " ".join([block[0] for block in line])

        # 添加段落
        doc.add_paragraph(paragraph_text)

    # 保存到临时文件
    temp_file_name = f"contract_{uuid.uuid4()}.docx"
    temp_path = os.path.join(UPLOAD_FOLDER, temp_file_name)
    doc.save(temp_path)

    return temp_path


def process_contract_image(image_bytes):
    """处理合同图片，返回Word文档路径"""
    try:
        # 将字节流转换为图片
        image = Image.open(BytesIO(image_bytes))
        img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

        # 执行OCR识别
        ocr_result = ocr.ocr(img, cls=True)

        # 根据OCR结果创建Word文档
        word_file_path = create_word_document_from_ocr(ocr_result)

        return word_file_path, None

    except Exception as e:
        return None, f"处理失败：{str(e)}"


# ======================== API接口 ========================
@app.route('/convert_contract', methods=['POST'])
def convert_contract_api():
    try:
        if 'image' not in request.files:
            return jsonify({"success": False, "error": "请上传图片文件"}), 400

        image_file = request.files['image']
        if image_file.filename == '':
            return jsonify({"success": False, "error": "未选择文件"}), 400

        image_bytes = image_file.read()
        word_path, error = process_contract_image(image_bytes)

        if error:
            return jsonify({"success": False, "error": error}), 400

        # 使用send_file发送生成的Word文件
        return send_file(
            word_path,
            as_attachment=True,
            download_name='converted_contract.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({"success": False, "error": f"服务器错误：{str(e)}"}), 500


# ======================== 启动服务 ========================
if __name__ == '__main__':
    print("正在初始化OCR模型并启动服务...")
    # 首次运行会下载模型，请耐心等待
    app.run(host='0.0.0.0', port=5001, debug=True)