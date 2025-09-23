from paddleocr import PaddleOCR
import numpy as np
import cv2
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor  # 新增RGBColor导入
from docx.oxml.ns import qn
import os
import uuid
import logging
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import sys
from pathlib import Path

# ======================== 配置与初始化 ========================
# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("ocr_service.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# 配置参数
class Config:
    # 输出文件存储配置
    OUTPUT_DOCS_DIR = Path("output_docs")
    # OCR配置
    OCR_LANGUAGE = 'ch'
    USE_ANGLE_CORRECTION = True
    LINE_HEIGHT_THRESHOLD = 20  # 行高阈值，用于文本块分组
    # Word文档字体配置
    DEFAULT_FONT_NAME = '微软雅黑'  # 统一使用的字体
    DEFAULT_FONT_SIZE = Pt(12)  # 统一使用的字体大小


config = Config()

# 确保输出目录存在
config.OUTPUT_DOCS_DIR.mkdir(parents=True, exist_ok=True)

# 初始化PaddleOCR
try:
    logger.info("正在初始化PaddleOCR模型...")
    ocr = PaddleOCR(
        use_angle_cls=config.USE_ANGLE_CORRECTION,
        lang=config.OCR_LANGUAGE,
        show_log=False  # 关闭PaddleOCR内部日志
    )
    logger.info("OCR模型初始化成功")
except Exception as e:
    logger.error(f"OCR模型初始化失败: {str(e)}", exc_info=True)
    messagebox.showerror("初始化失败", f"OCR模型初始化失败: {str(e)}")
    sys.exit(1)


# ======================== 工具函数 ========================
def validate_image_file(file_path):
    """验证文件是否为有效的图片"""
    allowed_extensions = {'png', 'jpg', 'jpeg', 'bmp', 'gif'}
    if '.' not in file_path:
        return False, "文件名格式不正确"

    ext = file_path.rsplit('.', 1)[1].lower()
    if ext not in allowed_extensions:
        return False, f"不支持的文件格式，支持的格式: {', '.join(allowed_extensions)}"

    # 检查文件是否存在
    if not os.path.exists(file_path):
        return False, "文件不存在"

    return True, "文件验证通过"


# ======================== 核心功能函数 ========================
def create_word_document_from_ocr(ocr_result):
    """
    根据OCR结果创建Word文档，统一字体和大小
    :param ocr_result: PaddleOCR返回的识别结果
    :return: 生成的Word文档的路径
    """
    try:
        doc = Document()

        # 设置文档默认字体
        style = doc.styles['Normal']
        style.font.name = config.DEFAULT_FONT_NAME
        style.font.size = config.DEFAULT_FONT_SIZE
        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT_NAME)

        # 添加标题和说明
        heading = doc.add_heading('合同图片转换结果', 0)
        heading.style.font.name = config.DEFAULT_FONT_NAME
        heading.style.font.size = Pt(16)
        heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT_NAME)

        # 添加说明段落并设置为红色
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run('注意：此文档由AI自动识别生成，可能存在误差，请仔细核对。')
        warning_run.font.color.rgb = RGBColor(255, 0, 0)  # 设置为红色
        # 应用其他字体样式
        warning_run.font.name = config.DEFAULT_FONT_NAME
        warning_run.font.size = config.DEFAULT_FONT_SIZE
        # 设置中文字体
        warning_run._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT_NAME)

        para = doc.add_paragraph('识别时间：' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        para.style = style

        para = doc.add_paragraph('--- 识别内容开始 ---')
        para.style = style

        if not ocr_result or not ocr_result[0]:
            para = doc.add_paragraph("未识别到任何文字内容。")
            para.style = style
            file_name = f"empty_{uuid.uuid4()}.docx"
            file_path = config.OUTPUT_DOCS_DIR / file_name
            doc.save(file_path)
            return str(file_path)

        # 提取文本块和其位置信息
        text_blocks = []
        for line in ocr_result[0]:
            poly, (text, score) = line
            x1, y1 = poly[0]
            x2, y2 = poly[2]
            text_blocks.append((text, int(x1), int(y1), int(x2), int(y2), score))

        # 布局还原核心逻辑
        # 1. 根据y坐标对文本块进行分组，模拟行
        lines = []
        current_line = []

        # 按y坐标排序
        text_blocks.sort(key=lambda b: b[2])

        for block in text_blocks:
            if not current_line:
                current_line.append(block)
            else:
                # 如果当前块的y坐标与行首块的y坐标差小于阈值，则认为在同一行
                if block[2] - current_line[0][2] < config.LINE_HEIGHT_THRESHOLD:
                    current_line.append(block)
                else:
                    lines.append(current_line)
                    current_line = [block]
        if current_line:
            lines.append(current_line)

        # 2. 逐行处理，写入Word文档
        for line in lines:
            # 同一行内，按x坐标从左到右排序
            line.sort(key=lambda b: b[1])

            # 将该行所有文本块合并成一个段落
            paragraph_text = " ".join([block[0] for block in line])

            # 添加段落并应用统一样式
            para = doc.add_paragraph(paragraph_text)
            para.style = style

        # 保存文件
        file_name = f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = config.OUTPUT_DOCS_DIR / file_name
        doc.save(file_path)

        logger.info(f"已生成Word文档: {file_path}")
        return str(file_path)

    except Exception as e:
        logger.error(f"创建Word文档失败: {str(e)}", exc_info=True)
        raise


def process_contract_image(image_path):
    """处理合同图片，返回Word文档路径"""
    try:
        # 打开图片文件
        with open(image_path, 'rb') as f:
            image_bytes = f.read()

        # 将字节流转换为图片
        image = Image.open(BytesIO(image_bytes))
        img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

        # 执行OCR识别
        logger.info("开始OCR识别...")
        ocr_result = ocr.ocr(img, cls=config.USE_ANGLE_CORRECTION)
        logger.info("OCR识别完成")

        # 根据OCR结果创建Word文档
        word_file_path = create_word_document_from_ocr(ocr_result)

        return word_file_path, None

    except Exception as e:
        error_msg = f"处理图片失败：{str(e)}"
        logger.error(error_msg, exc_info=True)
        return None, error_msg


# ======================== 界面相关函数 ========================
def select_image():
    """选择图片文件并处理"""
    file_path = filedialog.askopenfilename(
        title="选择图片文件",
        filetypes=[("图片文件", "*.png;*.jpg;*.jpeg;*.bmp;*.gif")]
    )

    if not file_path:
        return

    # 显示处理中提示
    status_label.config(text="正在处理，请稍候...")
    root.update()

    # 验证图片文件
    is_valid, msg = validate_image_file(file_path)
    if not is_valid:
        status_label.config(text="处理失败")
        messagebox.showerror("文件验证失败", msg)
        return

    # 处理图片
    word_path, error = process_contract_image(file_path)

    if error:
        status_label.config(text="处理失败")
        messagebox.showerror("处理失败", error)
    else:
        status_label.config(text="处理完成")
        messagebox.showinfo("处理完成", f"Word文档已生成：\n{word_path}")
        # 打开输出目录
        os.startfile(config.OUTPUT_DOCS_DIR)


# ======================== 主程序 ========================
if __name__ == '__main__':
    # 创建GUI窗口
    root = tk.Tk()
    root.title("纸质文书图片转Word工具")
    root.geometry("400x200")
    root.resizable(False, False)

    # 添加组件
    title_label = tk.Label(root, text="纸质文书图片转Word工具", font=("SimHei", 14))
    title_label.pack(pady=20)

    select_btn = tk.Button(root, text="选择图片文件", command=select_image, width=20, height=2, font=("SimHei", 10))
    select_btn.pack(pady=10)

    status_label = tk.Label(root, text="就绪", fg="gray", font=("SimHei", 10))
    status_label.pack(pady=20)

    # 启动主循环
    logger.info("程序启动成功，等待用户操作...")
    root.mainloop()
